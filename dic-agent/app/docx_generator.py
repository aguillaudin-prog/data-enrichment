"""Generate the DIC .docx file in either FRA short or ICAO long format.

We build the document programmatically with python-docx rather than from a static
template, because (a) the legs/segments are dynamic and (b) we want exact control
over the table grid that DIC ANNEX A requires.

Layout, numbering and field labels are aligned 1:1 with the reference DIC samples
in `dic-agent/sample-outputs/` (provided by the user). Any field, row or column
that doesn't appear in those references is removed; conversely, every block
present in the references is reproduced (state summary table per leg, IN CASE
OF EMERGENCY rows for diversion to alternate, RESERVED FOR ISSUING STATE block).
"""
from __future__ import annotations

import datetime as dt
from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

from app import db
from app.route_engine import LegResolution, format_zulu

GRAY = RGBColor(0xE0, 0xE0, 0xE0)
BLACK = RGBColor(0x00, 0x00, 0x00)


def _set_cell_bg(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _bold(cell, text: str, size: int = 9) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)


def _plain(cell, text: str, size: int = 9, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text or "")
    run.bold = bold
    run.font.size = Pt(size)


def _set_widths(table, widths_cm: list[float]) -> None:
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths_cm):
                cell.width = Cm(widths_cm[idx])


def _kv_row(table, label: str, value: str) -> None:
    row = table.add_row()
    _bold(row.cells[0], label)
    _plain(row.cells[1], value)


def _header_row(table, *labels: str, bg: str = "BFBFBF") -> None:
    row = table.add_row()
    for i, l in enumerate(labels):
        _bold(row.cells[i], l)
        _set_cell_bg(row.cells[i], bg)


_MONTHS_EN = {
    1: "JANUARY", 2: "FEBRUARY", 3: "MARCH", 4: "APRIL",
    5: "MAY", 6: "JUNE", 7: "JULY", 8: "AUGUST",
    9: "SEPTEMBER", 10: "OCTOBER", 11: "NOVEMBER", 12: "DECEMBER",
}


def _enrich_endpoint(label: str) -> str:
    """If `label` is the identifier of a user-added airport, append its DMS
    coordinates so the entry/exit cell in Appendix 1 unambiguously locates
    the aerodrome (reference DICs show the coords on those rows when the
    label has no published ICAO). Standard ICAO labels and border-crossing
    pseudo-labels (BORDER, EXIT BENIN, ENTRY NIGERIA) are returned as-is.
    """
    if not label:
        return label
    # Border-crossing pseudo-labels start with capital prefixes; never look
    # them up as airports.
    first_word = label.split()[0].upper()
    if first_word in ("BORDER", "ENTRY", "EXIT"):
        return label
    ap = db.find_airport(first_word)
    if not ap:
        return label
    try:
        if not ap["user_added"]:
            return label
    except (KeyError, IndexError):
        return label
    coords = f"{_format_dms_lat(ap['lat'])} / {_format_dms_lon(ap['lon'])}"
    return f"{label}\n{coords}"


def _format_dms_lat(lat: float) -> str:
    """Decimal latitude → 'N 9°34'45.56\"' (DMS, matching reference DIC notation)."""
    hemi = "N" if lat >= 0 else "S"
    lat = abs(lat)
    d = int(lat)
    m_full = (lat - d) * 60
    m = int(m_full)
    s = (m_full - m) * 60
    return f"{hemi} {d}°{m:02d}'{s:05.2f}\""


def _format_dms_lon(lon: float) -> str:
    hemi = "E" if lon >= 0 else "W"
    lon = abs(lon)
    d = int(lon)
    m_full = (lon - d) * 60
    m = int(m_full)
    s = (m_full - m) * 60
    return f"{hemi} {d}°{m:02d}'{s:05.2f}\""


def _format_airport(icao: str) -> str:
    """Render an airport for the DIC airport-list cells.

    Standard ICAO airports → 'CITY COUNTRY ICAO' (e.g. 'COTONOU BENIN DBBB').
    User-added operational labels (FOB, military AFB without a published
    ICAO) → '<NAME> <COUNTRY> <DMS coords>' — the 'ICAO' field IS the
    operational name, so no ICAO repeated, but the GPS coords appear so
    the receiving authority can locate the aerodrome unambiguously
    (matches reference DIC: 'TOUROU N 9°34'45.56" / E 3°14'7.09"').

    Returns the bare label if any lookup fails so we never lose the
    identifier in the rendered DIC.
    """
    icao = (icao or "").strip().upper()
    if not icao:
        return ""
    ap = db.find_airport(icao)
    if not ap:
        return icao
    is_user_added = False
    try:
        is_user_added = bool(ap["user_added"])
    except (KeyError, IndexError):
        is_user_added = False
    country = db.find_country_name(ap["country_iso"]) if ap["country_iso"] else None
    country_upper = (country or "").upper()
    name = (ap["name"] or "").strip()
    if is_user_added:
        # FOB / militaire non-publié : on garde le format avec coords pour
        # que l'autorité destinataire puisse géolocaliser sans ambiguïté.
        label = name.upper() if name and name.upper() != icao else icao
        coords = f"{_format_dms_lat(ap['lat'])} / {_format_dms_lon(ap['lon'])}"
        parts = [p for p in (label, country_upper, coords) if p]
        return " ".join(parts)
    # Standard ICAO airport : format opérateur de référence = 'CITY (ICAO)'
    # (avant : 'CITY COUNTRY ICAO' — la collègue ops a remonté que le format
    # avec parenthèses était le standard attendu, sans pays redondant car
    # déjà identifié par le préfixe ICAO 2-lettres).
    municipality = ""
    try:
        municipality = (ap["municipality"] or "").strip()
    except (KeyError, IndexError):
        municipality = ""
    if not municipality:
        municipality = name.split()[0] if name else ""
    return f"{municipality.upper()} ({icao})" if municipality else icao


def _format_airports_list(icaos: list[str]) -> str:
    return " / ".join(_format_airport(i) for i in icaos if (i or "").strip())


def _format_dic_time(t: dt.datetime | None) -> str:
    """Format like the reference: '04/05/2026 03H00Z'."""
    if not t:
        return ""
    return t.strftime("%d/%m/%Y %HH%MZ")


def _format_date_of_flight(eobts: list[dt.datetime]) -> str:
    """e.g. 'MAY 04 TO MAY 05, 2026' for first→last EOBT, or 'MAY 04, 2026'
    if the whole mission fits in a single day."""
    valid = [t for t in eobts if t]
    if not valid:
        return ""
    valid.sort()
    first, last = valid[0], valid[-1]
    if first.date() == last.date():
        return f"{_MONTHS_EN[first.month]} {first.day:02d}, {first.year}"
    return (
        f"{_MONTHS_EN[first.month]} {first.day:02d} TO "
        f"{_MONTHS_EN[last.month]} {last.day:02d}, {first.year}"
    )


def build_dic_document(mission: dict, leg_data: list[dict]) -> bytes:
    """Render the DIC docx in the unified Annex-A layout that matches the
    user's reference samples. The 'FRA' vs 'ICAO' distinction the codebase
    used to carry was retired: every reference DIC follows the same layout,
    so we ship a single format with the right field numbers, IN CASE OF
    EMERGENCY rows, etc. The `template_format` key on the mission dict is
    ignored."""
    return _build_dic(mission, leg_data)


def _build_dic(mission: dict, leg_data: list[dict]) -> bytes:
    """DIC docx, structured 1:1 with the reference samples in sample-outputs/.
    Layout, field numbers (1..40) and section headers come from the user's
    actual operational documents."""
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    # ── Title block ───────────────────────────────────────────────────
    # Two centred lines, matching the reference DIC header.
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p1.add_run("ANNEX A")
    r.bold = True
    r.font.size = Pt(11)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run("FRA DIPLOMATIC CLEARANCE (DIC) FORM")
    r.bold = True
    r.font.size = Pt(14)

    # ── (1) Reference + (2) Amendment on a SINGLE row of 4 cells,
    # then a second row for the Mission number — matches the reference
    # header layout exactly.
    t = doc.add_table(rows=0, cols=4)
    t.style = "Table Grid"
    r1 = t.add_row()
    _bold(r1.cells[0], "(1) Reference number :")
    _plain(r1.cells[1], mission.get("reference", ""))
    _bold(r1.cells[2], "(2) Amendment number :")
    _plain(r1.cells[3], mission.get("amendment", "V1"))
    r2 = t.add_row()
    _bold(r2.cells[0], "Mission number :")
    mc = r2.cells[1]
    mc.merge(r2.cells[2])
    mc.merge(r2.cells[3])
    _plain(mc, mission.get("mission_number", ""))
    _set_widths(t, [4.5, 4.5, 4.5, 4.5])
    doc.add_paragraph()

    # ── State summary table per leg (no VIP column, per reference) ───
    # Columns: STATE | R | N | L | DG | A | FR | EXISTING DIC NUMBER
    # = (3)..(10). One LEG header row, then one row per state crossed.
    state_tbl = doc.add_table(rows=0, cols=8)
    state_tbl.style = "Table Grid"
    hdr = state_tbl.add_row()
    _bold(hdr.cells[0], "(3) STATE")
    for i, lbl in enumerate(["(4) R", "(5) N", "(6) L", "(7) DG", "(8) A", "(9) FR", "(10) EXISTING DIC NUMBER"], start=1):
        _bold(hdr.cells[i], lbl)
        _set_cell_bg(hdr.cells[i], "BFBFBF")
    _set_cell_bg(hdr.cells[0], "BFBFBF")
    for li, leg in enumerate(leg_data, start=1):
        row = state_tbl.add_row()
        _bold(row.cells[0], f"LEG {li}")
        for i in range(1, 8):
            _plain(row.cells[i], "")
            _set_cell_bg(row.cells[i], "F2F2F2")
        for seg in leg["segments"]:
            row = state_tbl.add_row()
            _plain(row.cells[0], (seg["state_name"] or "").upper())
            _plain(row.cells[1], "X" if seg.get("R") else "")
            _plain(row.cells[2], "X" if seg.get("N") else "")
            _plain(row.cells[3], "X" if seg.get("L") else "")
            _plain(row.cells[4], "X" if seg.get("DG") else "")
            _plain(row.cells[5], "X" if seg.get("A") else "")
            _plain(row.cells[6], seg.get("FR", "I/V"))
            _plain(row.cells[7], seg.get("existing_dic", ""))
    _set_widths(state_tbl, [3.5, 1.0, 1.0, 1.0, 1.0, 1.0, 1.2, 4.3])
    doc.add_paragraph()

    # ── Info table — sections (11..34) ───────────────────────────────
    info_tbl = doc.add_table(rows=0, cols=3)
    info_tbl.style = "Table Grid"
    _header_row(info_tbl, "SERIAL", "REQUESTED INFORMATION", "INFORMATION SUBMITTED")

    def _info(serial: str, label: str, value: str) -> None:
        r = info_tbl.add_row()
        _plain(r.cells[0], serial)
        _plain(r.cells[1], label)
        _plain(r.cells[2], value)

    def _section(label: str) -> None:
        r = info_tbl.add_row()
        for c in r.cells:
            _set_cell_bg(c, "BFBFBF")
        _bold(r.cells[1], label)

    # Aircraft and crew (11..20)
    _section("AIRCRAFT AND CREW")
    _info("(11)", "Requesting state", mission.get("requesting_state", "FRANCE"))
    _info("(11a)", "Operator", mission.get("operator", ""))
    # Number and type of aircraft: '01 DHC6-400 TY-BAB' format
    ac_type = (mission.get("aircraft_type_icao") or "").strip()
    reg = (mission.get("registration") or "").strip()
    nb_ac = "01"
    aircraft_str = " ".join(p for p in [nb_ac, ac_type, reg] if p)
    _info("(12)", "Number and type of aircraft", aircraft_str)
    _info("(13)", "Aircraft registration", reg)
    _info("(14)", "Spare aircraft", mission.get("spare_aircraft", "/"))
    _info("(15)", "Callsign (including spare if different)", mission.get("callsign", ""))
    n_crew = mission.get("n_crew", 2)
    try:
        n_crew_str = f"{int(n_crew):02d}"
    except (TypeError, ValueError):
        n_crew_str = str(n_crew or "")
    _info("(16)", "Number of crew members", n_crew_str)
    _info("(17)", "Pilot rank and name", mission.get("pilots", ""))
    # (18)..(20) Photographic sensors / Armament / EW — removed per user
    # request. If a future mission needs to disclose them, the user adds
    # the info to the Remarks (30) field instead.

    # Flight details (21..26)
    _section("FLIGHT DETAILS (Detailed routing in Appendix 1)")
    eobts = [leg.get("eobt") for leg in leg_data if leg.get("eobt")]
    _info("(21)", "Date of flight", mission.get("date_of_flight") or _format_date_of_flight(eobts))
    _info("(22)", "Purpose of flight", mission.get("purpose", ""))
    departures = [leg.get("origin") for leg in leg_data if leg.get("origin")]
    destinations = [leg.get("destination") for leg in leg_data if leg.get("destination")]
    alternates = [leg.get("alternate") for leg in leg_data if leg.get("alternate")]
    _info("(23)", "Departure airport(s)", _format_airports_list(departures))
    _info("(24)", "Destination airport(s)", _format_airports_list(destinations))
    _info("(25)", "Alternate airport(s)", _format_airports_list(alternates))
    _info("(26)", "Radio frequencies", mission.get("radio_frequencies") or "V/U/HF")

    # Load information (27..29) — only the passenger count is exposed.
    # (28) VIP and (29) DG removed per user request: when a flight actually
    # carries a VIP or DG, the operator surfaces it in the Remarks block.
    _section("LOAD INFORMATION")
    _info("(27)", "Number of passengers", mission.get("n_passengers", "TBN"))

    # Remarks (30)
    _section("REMARKS")
    _info("(30)", "", mission.get("remarks", ""))

    # Point of Contact (31..34)
    _section("POINT OF CONTACT")
    _info("(31)", "Rank, name, first name", mission.get("poc_name", ""))
    _info("(32)", "Telephone number", mission.get("poc_phone", ""))
    _info("(33)", "E-mail", mission.get("poc_email_functional", ""))
    _info("(34)", "Fax", mission.get("poc_fax", ""))

    # Reserved for issuing state (35..36) + stamp/date/signature rows —
    # empty placeholders that the receiving authority fills in.
    _section("RESERVED FOR ISSUING STATE")
    _info("(35)", "STATE ISSUING", "")
    _info("(36)", "DIPLOMATIC CLEARANCE NUMBER & VALIDITY", "")

    _set_widths(info_tbl, [1.5, 6, 9.5])
    doc.add_paragraph()

    stamp_tbl = doc.add_table(rows=0, cols=2)
    stamp_tbl.style = "Table Grid"
    r1 = stamp_tbl.add_row()
    _bold(r1.cells[0], "Stamp issuing state :")
    _bold(r1.cells[1], "Date :")
    r2 = stamp_tbl.add_row()
    _plain(r2.cells[0], "")
    _plain(r2.cells[1], "")
    r3 = stamp_tbl.add_row()
    _bold(r3.cells[0], "Signature :")
    sc = r3.cells[0]
    sc.merge(r3.cells[1])
    _set_widths(stamp_tbl, [8.5, 8.5])

    doc.add_page_break()

    # ── Appendix 1 — Detailed itinerary per leg ──────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ANNEX A, APPENDIX 1")
    r.bold = True
    r.font.size = Pt(12)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run("DETAILED ITINERARY")
    r.bold = True
    r.font.size = Pt(11)

    callsign_for_legs = mission.get("callsign", "")

    for li, leg in enumerate(leg_data, start=1):
        head = doc.add_paragraph()
        leg_label = (
            f"Leg {li}    "
            f"{leg.get('origin','')} → {leg.get('destination','')}    "
            f"Callsign : {callsign_for_legs}"
        )
        run = head.add_run(leg_label)
        run.bold = True
        run.font.size = Pt(10)

        # 6-column itinerary table: State | Entry | Route | Exit | FL | TAS
        # Matches reference Appendix 1 header (39)..(44).
        leg_tbl = doc.add_table(rows=0, cols=6)
        leg_tbl.style = "Table Grid"
        hdr = leg_tbl.add_row()
        _bold(hdr.cells[0], "(39) State")
        _bold(hdr.cells[1], "(40) Entry point and timing or airfield + EOBT")
        _bold(hdr.cells[2], "(41) Route over Territory")
        _bold(hdr.cells[3], "(42) Exit point and timing or airfield + EIBT")
        _bold(hdr.cells[4], "(43) FL")
        _bold(hdr.cells[5], "(44) TAS")
        for c in hdr.cells:
            _set_cell_bg(c, "BFBFBF")
        for seg in leg["segments"]:
            row = leg_tbl.add_row()
            _plain(row.cells[0], (seg["state_name"] or "").upper())
            _plain(row.cells[1], f"{_enrich_endpoint(seg['entry_label'])}\n{seg['entry_time_str']}")
            _plain(row.cells[2], seg["route_in_country"])
            _plain(row.cells[3], f"{_enrich_endpoint(seg['exit_label'])}\n{seg['exit_time_str']}")
            _plain(row.cells[4], str(seg.get("fl") or ""))
            _plain(row.cells[5], str(seg.get("tas") or ""))

        # IN CASE OF EMERGENCY — diversion to the alternate.
        # Reference layout (4 columns, italicised):
        #   row N+0:   "IN CASE OF EMERGENCY"  spanning full width (label)
        #   row N+1:   <country>  |  (empty)  |  DCT  |  <ICAO> <ETA>
        # We don't have the navaid name near the alternate (would require a
        # mapping table), so the route cell shows just 'DCT'. The ETA is the
        # leg's destination ETA + 30 min as a rough estimate.
        alt = (leg.get("alternate") or "").strip().upper()
        if alt:
            label_row = leg_tbl.add_row()
            for c in label_row.cells:
                _set_cell_bg(c, "FFF2CC")
            label_cell = label_row.cells[0]
            for ci in range(1, 6):
                label_cell.merge(label_row.cells[ci])
            _bold(label_cell, "IN CASE OF EMERGENCY")

            alt_ap = db.find_airport(alt)
            alt_country = (
                db.find_country_name(alt_ap["country_iso"])
                if alt_ap and alt_ap["country_iso"] else ""
            )
            eta_str = ""
            fl_str = ""
            tas_str = ""
            if leg.get("segments"):
                last_seg = leg["segments"][-1]
                eta_str = last_seg.get("exit_time_str") or ""
                fl_str = str(last_seg.get("fl") or "")
                tas_str = str(last_seg.get("tas") or "")
            row = leg_tbl.add_row()
            for c in row.cells:
                _set_cell_bg(c, "FFF2CC")
            _plain(row.cells[0], (alt_country or "—").upper())
            _plain(row.cells[1], "")
            _plain(row.cells[2], "DCT")
            _plain(row.cells[3], f"{alt}  {eta_str}".strip())
            _plain(row.cells[4], fl_str)
            _plain(row.cells[5], tas_str)
        _set_widths(leg_tbl, [2.5, 4.0, 4.5, 4.0, 1.0, 1.0])
        doc.add_paragraph()

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def serialize_leg(leg_input: dict, resolution: LegResolution, style: str = "FRA") -> dict:
    segs = []
    for seg in resolution.segments:
        entry_t = format_zulu(seg.entry_time, style) if seg.entry_time else ""
        exit_t = format_zulu(seg.exit_time, style) if seg.exit_time else ""
        is_origin_country = seg.state_iso and seg.state_iso == leg_input.get("origin_iso")
        is_dest_country = seg.state_iso and seg.state_iso == leg_input.get("destination_iso")
        overrides = leg_input.get("overrides", {}).get(seg.state_iso, {})
        segs.append(
            {
                "state_name": overrides.get("state_name", seg.state_name),
                "state_iso": seg.state_iso,
                "entry_label": seg.entry_label,
                "exit_label": seg.exit_label,
                "entry_time_str": entry_t,
                "exit_time_str": exit_t,
                "route_in_country": seg.route_in_country,
                "fl": seg.fl,
                "tas": seg.tas,
                "R": overrides.get("R", is_origin_country or is_dest_country),
                "N": overrides.get("N", not (is_origin_country or is_dest_country)),
                "L": overrides.get("L", is_origin_country or is_dest_country),
                "VIP": overrides.get("VIP", False),
                "DG": overrides.get("DG", False),
                "A": overrides.get("A", False),
                "FR": overrides.get("FR", "I"),
                "existing_dic": overrides.get("existing_dic", ""),
            }
        )
    return {
        "origin": leg_input.get("origin", ""),
        "destination": leg_input.get("destination", ""),
        "alternate": leg_input.get("alternate", ""),
        "eobt": leg_input.get("eobt"),
        "segments": segs,
    }
